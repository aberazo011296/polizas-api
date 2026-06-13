"""
Tests de construcción de templates (services/template_builder.py).
"""
import io

import pytest
from docx import Document

from app.services.template_builder import construir_template


def _docx(parrafos: list[str]) -> bytes:
    doc = Document()
    for p in parrafos:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _texto(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


class TestConstruirTemplate:
    def test_reemplaza_texto_por_marcador(self):
        original = _docx(["El asegurado Juan Perez tiene cobertura."])
        salida = construir_template(original, {"nombre_asegurado": "Juan Perez"})
        texto = _texto(salida)
        assert "{{nombre_asegurado}}" in texto
        assert "Juan Perez" not in texto

    def test_reemplazo_vacio_se_ignora(self):
        original = _docx(["Texto sin cambios."])
        salida = construir_template(original, {"var": "   "})
        # Un reemplazo vacío no debe introducir marcadores
        assert "{{var}}" not in _texto(salida)

    def test_multiples_reemplazos(self):
        original = _docx(["Poliza 990664 a nombre de Ana Lopez."])
        salida = construir_template(
            original,
            {"numero_poliza": "990664", "nombre_asegurado": "Ana Lopez"},
        )
        texto = _texto(salida)
        assert "{{numero_poliza}}" in texto
        assert "{{nombre_asegurado}}" in texto

    def test_tolerante_a_espacios(self):
        # El texto del Word puede tener espacios distintos al buscado
        original = _docx(["El  asegurado   Juan."])
        salida = construir_template(original, {"nombre": "El asegurado Juan."})
        assert "{{nombre}}" in _texto(salida)
