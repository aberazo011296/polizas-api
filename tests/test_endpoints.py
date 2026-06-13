"""
Tests de endpoints (routers): plantillas, templates, certificados, pólizas.

Cubren el flujo completo: crear plantilla → subir template → generar
certificado, además de validaciones (404, 400, 413, 503).
"""
import io
import json

import pytest
from docx import Document

from app.core.config import settings


def _docx(parrafos: list[str]) -> bytes:
    doc = Document()
    for p in parrafos:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestPlantillaCRUD:
    def test_crear_solo_variables(self, plantilla_creada):
        assert plantilla_creada["id"]
        assert len(plantilla_creada["variables"]) == 2

    def test_crear_sin_definicion_falla(self, client):
        resp = client.post("/plantillas", json={
            "nombre": "Vacia", "aseguradora": "a", "tipo_poliza": "b",
        })
        assert resp.status_code == 400

    def test_obtener(self, client, plantilla_creada):
        resp = client.get(f"/plantillas/{plantilla_creada['id']}")
        assert resp.status_code == 200
        assert resp.json()["nombre"] == "Fixture Plantilla"

    def test_actualizar(self, client, plantilla_creada):
        body = {
            "nombre": "Renombrada",
            "aseguradora": "testaseg",
            "tipo_poliza": "testtipo",
            "variables": [{"nombre": "x", "descripcion": ""}],
        }
        resp = client.put(f"/plantillas/{plantilla_creada['id']}", json=body)
        assert resp.status_code == 200
        assert resp.json()["nombre"] == "Renombrada"

    def test_actualizar_no_existe(self, client):
        body = {"nombre": "x", "aseguradora": "a", "tipo_poliza": "b",
                "variables": [{"nombre": "v", "descripcion": ""}]}
        resp = client.put("/plantillas/no-existe", json=body)
        assert resp.status_code == 404


class TestTemplateUpload:
    def test_subir_descargar_template(self, client, plantilla_creada):
        pid = plantilla_creada["id"]
        docx = _docx(["Asegurado: {{nombre_asegurado}}"])
        resp = client.post(
            f"/plantillas/{pid}/template",
            files={"archivo": ("t.docx", docx, "application/octet-stream")},
        )
        assert resp.status_code == 200
        # Ahora se puede descargar
        resp2 = client.get(f"/plantillas/{pid}/template")
        assert resp2.status_code == 200

    def test_subir_no_docx_rechaza(self, client, plantilla_creada):
        pid = plantilla_creada["id"]
        resp = client.post(
            f"/plantillas/{pid}/template",
            files={"archivo": ("t.txt", b"hola", "text/plain")},
        )
        assert resp.status_code == 400

    def test_subir_template_excede_tamano(self, client, plantilla_creada, monkeypatch):
        monkeypatch.setattr(settings, "max_docx_size_bytes", 100)
        pid = plantilla_creada["id"]
        grande = _docx(["x" * 5000])
        resp = client.post(
            f"/plantillas/{pid}/template",
            files={"archivo": ("t.docx", grande, "application/octet-stream")},
        )
        assert resp.status_code == 413

    def test_descargar_template_inexistente(self, client, plantilla_creada):
        # Plantilla existe pero nunca subió template
        resp = client.get(f"/plantillas/{plantilla_creada['id']}/template")
        assert resp.status_code == 404

    def test_build_excede_tamano(self, client, plantilla_creada, monkeypatch):
        monkeypatch.setattr(settings, "max_docx_size_bytes", 100)
        pid = plantilla_creada["id"]
        grande = _docx(["x" * 5000])
        resp = client.post(
            f"/plantillas/{pid}/template/build",
            files={"archivo": ("t.docx", grande, "application/octet-stream")},
            data={"reemplazos": json.dumps({"nombre_asegurado": "x" * 5000})},
        )
        assert resp.status_code == 413

    def test_doc_referencia_inexistente(self, client, plantilla_creada):
        resp = client.get(f"/plantillas/{plantilla_creada['id']}/template/doc-referencia")
        assert resp.status_code == 404


class TestCertificadoFlujoCompleto:
    def test_generar_certificado_end_to_end(self, client, plantilla_creada):
        pid = plantilla_creada["id"]
        # 1. Subir template con marcadores
        docx = _docx(["Asegurado: {{nombre_asegurado}}", "Póliza: {{numero_poliza}}"])
        up = client.post(
            f"/plantillas/{pid}/template",
            files={"archivo": ("t.docx", docx, "application/octet-stream")},
        )
        assert up.status_code == 200
        # 2. Generar certificado
        resp = client.post("/certificados/generar", json={
            "plantilla_id": pid,
            "variables": {"nombre_asegurado": "Maria Vaca", "numero_poliza": "990633"},
        })
        assert resp.status_code == 200
        assert resp.headers["content-type"].startswith(
            "application/vnd.openxmlformats")
        doc = Document(io.BytesIO(resp.content))
        texto = "\n".join(p.text for p in doc.paragraphs)
        assert "Maria Vaca" in texto
        assert "990633" in texto

    def test_generar_plantilla_no_existe(self, client):
        resp = client.post("/certificados/generar", json={
            "plantilla_id": "no-existe", "variables": {},
        })
        assert resp.status_code == 404

    def test_generar_sin_template_subido(self, client, plantilla_creada):
        # Plantilla existe pero no tiene template .docx → 404 de generación
        resp = client.post("/certificados/generar", json={
            "plantilla_id": plantilla_creada["id"], "variables": {"x": "y"},
        })
        assert resp.status_code == 404


class TestPolizasUpload:
    def test_upload_tipo_invalido(self, client, plantilla_creada):
        resp = client.post("/polizas/upload", files={
            "archivo": ("t.txt", b"hola", "text/plain"),
        }, data={"plantilla_id": plantilla_creada["id"]})
        assert resp.status_code == 415

    def test_upload_vacio(self, client, plantilla_creada):
        resp = client.post("/polizas/upload", files={
            "archivo": ("t.pdf", b"", "application/pdf"),
        }, data={"plantilla_id": plantilla_creada["id"]})
        assert resp.status_code == 400

    def test_upload_plantilla_no_existe(self, client, pdf_bytes):
        resp = client.post("/polizas/upload", files={
            "archivo": ("t.pdf", pdf_bytes, "application/pdf"),
        }, data={"plantilla_id": "no-existe"})
        assert resp.status_code == 404

    def test_upload_excede_tamano(self, client, plantilla_creada, monkeypatch):
        monkeypatch.setattr(settings, "max_pdf_size_bytes", 10)
        resp = client.post("/polizas/upload", files={
            "archivo": ("t.pdf", b"%PDF-1.4" + b"x" * 100, "application/pdf"),
        }, data={"plantilla_id": plantilla_creada["id"]})
        assert resp.status_code == 413

    def test_upload_sin_anthropic_ni_cajas_502(self, client, plantilla_creada, pdf_bytes, monkeypatch):
        # Sin IA y plantilla solo-variables (sin cajas) → 502
        monkeypatch.setattr(settings, "anthropic_api_key", None)
        resp = client.post("/polizas/upload", files={
            "archivo": ("t.pdf", pdf_bytes, "application/pdf"),
        }, data={"plantilla_id": plantilla_creada["id"]})
        assert resp.status_code == 502


class TestSugerirVariables:
    def test_sin_anthropic_key_503(self, client, pdf_bytes, monkeypatch):
        monkeypatch.setattr(settings, "anthropic_api_key", None)
        resp = client.post("/plantillas/sugerir-variables", files={
            "archivo": ("t.pdf", pdf_bytes, "application/pdf"),
        })
        assert resp.status_code == 503

    def test_tipo_invalido_415(self, client):
        resp = client.post("/plantillas/sugerir-variables", files={
            "archivo": ("t.txt", b"hola", "text/plain"),
        })
        assert resp.status_code == 415
