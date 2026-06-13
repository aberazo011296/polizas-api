"""
Fixtures compartidas para la suite de tests.
"""
import io
import os

# Forzar backend filesystem para los tests, sin importar STORAGE_BACKEND en
# el .env del desarrollador (p.ej. si está en "mongo" para correr localmente).
# El selector de app/storage/__init__.py resuelve a nivel de import, así que
# debe fijarse antes de importar app.main. Los tests de Mongo (mongomock)
# trabajan directo contra app.storage.mongo, sin pasar por este selector.
os.environ.setdefault("STORAGE_BACKEND", "local")

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.core.config import settings
from app.core.paths import ruta_template_docx
from app.main import app
from app.storage.local import eliminar_plantilla


@pytest.fixture
def client():
    return TestClient(app)


@pytest.fixture
def pdf_bytes():
    """PDF mínimo de una página en blanco (PyMuPDF)."""
    import fitz
    doc = fitz.open()
    doc.new_page()
    return doc.tobytes()


def construir_docx(parrafos: list[str]) -> bytes:
    """Crea un .docx en memoria con los párrafos dados."""
    doc = Document()
    for p in parrafos:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def texto_de_docx(docx_bytes: bytes) -> str:
    """Extrae todo el texto de un .docx (para asserts)."""
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


@pytest.fixture
def make_docx():
    """Factory de .docx para los tests."""
    return construir_docx


@pytest.fixture
def docx_text():
    """Helper para leer el texto de un .docx generado."""
    return texto_de_docx


@pytest.fixture
def plantilla_creada(client):
    """
    Crea una plantilla de prueba y la elimina al terminar.
    Devuelve el dict de la plantilla creada.
    """
    body = {
        "nombre": "Fixture Plantilla",
        "aseguradora": "testaseg",
        "tipo_poliza": "testtipo",
        "variables": [
            {"nombre": "nombre_asegurado", "descripcion": "nombre del asegurado"},
            {"nombre": "numero_poliza", "descripcion": "número de póliza"},
        ],
    }
    resp = client.post("/plantillas", json=body)
    assert resp.status_code == 201
    data = resp.json()
    yield data
    # Limpieza: plantilla, su template .docx y el doc de referencia, para
    # no contaminar tests que esperan "sin template subido".
    try:
        eliminar_plantilla(data["id"])
    except Exception:
        pass
    try:
        ruta = ruta_template_docx(data["aseguradora"], data["tipo_poliza"])
        ruta.unlink(missing_ok=True)
    except Exception:
        pass
    try:
        (settings.data_dir / "referencias" / f"{data['id']}.docx").unlink(missing_ok=True)
    except Exception:
        pass
