"""
Tests de generación de certificados (services/generador.py).
"""
import io

import pytest
from docx import Document

from app.core.config import settings
from app.core.errors import PlantillaNoEncontradaError
from app.core.paths import ruta_template_docx
from app.services.generador import generar_certificado


def _crear_template_docx(aseguradora: str, tipo: str, parrafos: list[str]) -> None:
    """Escribe un template .docx con marcadores Jinja en templates_dir."""
    ruta = ruta_template_docx(aseguradora, tipo)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for p in parrafos:
        doc.add_paragraph(p)
    doc.save(str(ruta))


def _texto(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


@pytest.fixture
def template_simple():
    aseg, tipo = "genaseg", "gentipo"
    _crear_template_docx(aseg, tipo, ["Asegurado: {{nombre_asegurado}}",
                                      "Póliza: {{numero_poliza}}"])
    yield aseg, tipo
    ruta = ruta_template_docx(aseg, tipo)
    if ruta.exists():
        ruta.unlink()


class TestGenerarCertificado:
    def test_rellena_variables(self, template_simple):
        aseg, tipo = template_simple
        docx_bytes, usadas, faltantes = generar_certificado(
            aseg, tipo, {"nombre_asegurado": "Juan Perez", "numero_poliza": "990664"}
        )
        texto = _texto(docx_bytes)
        assert "Juan Perez" in texto
        assert "990664" in texto
        assert set(usadas) == {"nombre_asegurado", "numero_poliza"}
        assert faltantes == []

    def test_variable_faltante(self, template_simple):
        aseg, tipo = template_simple
        docx_bytes, usadas, faltantes = generar_certificado(
            aseg, tipo, {"nombre_asegurado": "Ana"}
        )
        assert "numero_poliza" in faltantes
        assert "nombre_asegurado" in usadas

    def test_template_no_existe(self):
        with pytest.raises(PlantillaNoEncontradaError):
            generar_certificado("noexiste", "notipo", {})

    def test_salto_de_linea_genera_parrafos(self):
        aseg, tipo = "braseg", "brtipo"
        _crear_template_docx(aseg, tipo, ["Coberturas: {{listado}}"])
        try:
            docx_bytes, _, _ = generar_certificado(
                aseg, tipo, {"listado": "a) Muerte\nb) Invalidez"}
            )
            doc = Document(io.BytesIO(docx_bytes))
            textos = [p.text for p in doc.paragraphs]
            # El \n debe haber generado párrafos separados
            assert any("Muerte" in t for t in textos)
            assert any("Invalidez" in t for t in textos)
        finally:
            ruta = ruta_template_docx(aseg, tipo)
            if ruta.exists():
                ruta.unlink()


class TestCoberturasLoop:
    """Grupo repetible de coberturas: {% for c in coberturas %}."""

    def _template_loop(self, aseg, tipo):
        _crear_template_docx(aseg, tipo, [
            "Póliza: {{numero_poliza}}",
            "{% for c in coberturas %}",
            "Cobertura: {{c.nombre}} — {{c.suma_asegurada}}",
            "Exclusiones:{{c.listado_exclusiones}}",
            "{% endfor %}",
        ])

    def test_genera_una_seccion_por_cobertura(self):
        aseg, tipo = "cobaseg", "cobtipo"
        self._template_loop(aseg, tipo)
        try:
            coberturas = [
                {"nombre": "Muerte por cualquier causa", "suma_asegurada": "200,000.00",
                 "listado_exclusiones": "a) Suicidio\nb) Guerra"},
                {"nombre": "Anticipo Enfermedades Graves", "suma_asegurada": "50%",
                 "listado_exclusiones": "a) Preexistencias"},
            ]
            docx_bytes, usadas, faltantes = generar_certificado(
                aseg, tipo, {"numero_poliza": "990614"}, coberturas
            )
            texto = _texto(docx_bytes)
            assert "Muerte por cualquier causa" in texto
            assert "Anticipo Enfermedades Graves" in texto
            assert "50%" in texto
            # El \n dentro del loop también se expande a párrafos
            assert "Suicidio" in texto and "Guerra" in texto
            # `coberturas` no debe contarse como variable plana faltante
            assert "coberturas" not in faltantes
            assert "numero_poliza" in usadas
        finally:
            ruta = ruta_template_docx(aseg, tipo)
            if ruta.exists():
                ruta.unlink()

    def test_sin_coberturas_no_rompe(self):
        # Plantilla con loop pero póliza de 0 coberturas: el loop no imprime
        # nada y la generación no falla.
        aseg, tipo = "cobaseg2", "cobtipo2"
        self._template_loop(aseg, tipo)
        try:
            docx_bytes, _, faltantes = generar_certificado(
                aseg, tipo, {"numero_poliza": "990000"}, []
            )
            texto = _texto(docx_bytes)
            assert "990000" in texto
            assert "coberturas" not in faltantes
        finally:
            ruta = ruta_template_docx(aseg, tipo)
            if ruta.exists():
                ruta.unlink()
